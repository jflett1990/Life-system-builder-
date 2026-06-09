"""
Tests for DOCX export — endpoint headers + bytes, and DocxBuilder content.

Run with:
  cd artifacts/api-server && python -m pytest tests/test_docx_export.py -v
"""
from __future__ import annotations

import io
import sys
import os

import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))


# ── DocxBuilder unit tests (no server required) ───────────────────────────────

MINIMAL_OUTPUTS: dict = {
    "system_architecture": {
        "system_name": "Build a CRUD App with Supabase",
        "topic": "Build a CRUD app with Supabase",
        "system_objective": "Finish with a working CRUD app talking to a Supabase Postgres table.",
        "time_horizon": "3 hours hands-on",
        "audience": "Self-taught developer",
    },
    "chapter_expansion": {
        "chapters": [
            {
                "chapter_number": 1,
                "chapter_title": "Project Setup",
                "narrative": "This chapter scaffolds the project and verifies the toolchain.",
                "quick_reference_rules": ["Rule one.", "Rule two."],
                "worksheets": [
                    {
                        "id": "ws-form",
                        "title": "Project Config Record",
                        "purpose": "Record the project's configuration values for later steps.",
                        "layout": "form",
                        "estimated_completion_time": "10 minutes",
                        "sections": [
                            {
                                "section_title": "Core Settings",
                                "instructions": "Fill in all fields from your terminal output.",
                                "fields": [
                                    {"label": "Project Name", "type": "text", "placeholder": "e.g. supabase-crud-app"},
                                    {"label": "Primary Framework", "type": "text"},
                                    {
                                        "label": "Project Stage",
                                        "type": "select",
                                        "options": ["Scaffold", "CRUD Working", "Deployed"],
                                    },
                                ],
                            }
                        ],
                        "decision_gates": [
                            {
                                "gate_id": "gate-1",
                                "condition": "Does `npm run dev` start without errors?",
                                "pass_action": "Proceed to database setup.",
                                "fail_action": "Fix install errors before continuing.",
                            }
                        ],
                    },
                    {
                        "id": "ws-checklist",
                        "title": "Setup Verification Checklist",
                        "purpose": "Verify the environment is ready before building.",
                        "layout": "checklist",
                        "checklist_items": [
                            "Verify node --version is 18.17 or newer",
                            "Create a Supabase project in the dashboard",
                            "Copy the API URL and anon key into .env.local",
                        ],
                        "decision_gates": [],
                    },
                    {
                        "id": "ws-table",
                        "title": "Command Reference",
                        "purpose": "Track the commands used in this step and their expected output.",
                        "layout": "table",
                        "table_columns": ["Command", "Run From", "Expected Output", "Common Failure"],
                        "table_row_count": 8,
                        "decision_gates": [],
                    },
                    {
                        "id": "ws-two-col",
                        "title": "Local vs Production Config",
                        "purpose": "Map local configuration to its production equivalent.",
                        "layout": "two-column",
                        "left_column_label": "Local",
                        "right_column_label": "Production",
                        "sections": [
                            {
                                "section_title": "Environment",
                                "fields": [
                                    {"label": "API URL", "type": "text"},
                                    {"label": "Anon Key Source", "type": "text"},
                                ],
                            }
                        ],
                        "decision_gates": [],
                    },
                ],
            }
        ]
    },
}


def _build_doc(outputs: dict | None = None) -> bytes:
    from render.docx_builder import DocxBuilder

    return DocxBuilder().build(project_id=1, all_outputs=outputs or MINIMAL_OUTPUTS)


class TestDocxBuilderBytes:
    def test_returns_non_empty_bytes(self) -> None:
        result = _build_doc()
        assert isinstance(result, bytes)
        assert len(result) > 1_000, f"Expected > 1 KB, got {len(result)} bytes"

    def test_starts_with_zip_magic(self) -> None:
        result = _build_doc()
        assert result[:2] == b"PK", "DOCX must start with ZIP PK magic bytes"


class TestDocxBuilderStructure:
    """Parse the generated DOCX and assert heading hierarchy + content."""

    def _parse(self, outputs: dict | None = None):
        from docx import Document

        raw = _build_doc(outputs)
        return Document(io.BytesIO(raw))

    def test_heading_hierarchy(self) -> None:
        doc = self._parse()
        paragraphs = doc.paragraphs
        h1 = [p.text for p in paragraphs if p.style.name == "Heading 1"]
        h2 = [p.text for p in paragraphs if p.style.name == "Heading 2"]
        h3 = [p.text for p in paragraphs if p.style.name == "Heading 3"]

        assert any("Chapter 1" in t for t in h1), f"Expected chapter H1, got: {h1}"
        assert any("Project Config Record" in t for t in h2), f"H2 missing worksheet: {h2}"
        assert any("Project Name" in t for t in h3), f"H3 field label missing: {h3}"

    def test_form_layout_fill_lines(self) -> None:
        doc = self._parse()
        fill_lines = [p for p in doc.paragraphs if p.text.startswith("___")]
        assert len(fill_lines) >= 2, f"Expected fill-in lines, got {len(fill_lines)}"

    def test_form_layout_select_options(self) -> None:
        doc = self._parse()
        option_texts = [p.text for p in doc.paragraphs if "○" in p.text]
        assert len(option_texts) >= 1, "Expected radio-circle option items for 'select' field"

    def test_checklist_layout_checkboxes(self) -> None:
        doc = self._parse()
        checkbox_paras = [p for p in doc.paragraphs if "☐" in p.text]
        assert len(checkbox_paras) >= 3, f"Expected ≥3 checkbox items, got {len(checkbox_paras)}"

    def test_table_layout_word_table(self) -> None:
        doc = self._parse()
        assert len(doc.tables) >= 1, "Expected at least one Word table for 'table' layout"
        headers = [c.text for c in doc.tables[0].rows[0].cells]
        assert "Command" in headers, f"Expected 'Command' column, got: {headers}"

    def test_two_column_word_table(self) -> None:
        doc = self._parse()
        two_col_tables = [t for t in doc.tables if len(t.columns) == 3]
        assert len(two_col_tables) >= 1, "Expected a 3-column Word table for 'two-column' layout"
        headers = [c.text for c in two_col_tables[0].rows[0].cells]
        assert "Local" in headers, f"Expected 'Local' header, got: {headers}"
        assert "Production" in headers, f"Expected 'Production' header, got: {headers}"

    def test_decision_gates_present(self) -> None:
        doc = self._parse()
        gate_paras = [p for p in doc.paragraphs if "Gate:" in p.text]
        assert len(gate_paras) >= 1, "Expected decision gate paragraph"

    def test_cover_content(self) -> None:
        doc = self._parse()
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Build a CRUD App with Supabase" in full_text
        assert "Build a CRUD app with Supabase" in full_text

    def test_fallback_no_chapters(self) -> None:
        outputs = {
            "system_architecture": {"system_name": "Fallback Test"},
            "worksheet_system": {
                "worksheets": [
                    {
                        "id": "ws-1",
                        "title": "Legacy Worksheet",
                        "purpose": "A worksheet from the old stage.",
                        "sections": [
                            {"section_title": "Info", "fields": [{"label": "Name", "type": "text"}]}
                        ],
                    }
                ]
            },
        }
        doc = self._parse(outputs)
        h1 = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
        assert any("Legacy Worksheet" in t for t in h1), f"Expected legacy worksheet as H1: {h1}"
