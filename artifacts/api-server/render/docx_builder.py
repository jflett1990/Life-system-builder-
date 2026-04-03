"""
DocxBuilder — Generates a Word (.docx) document from Life System Builder pipeline data.

Document structure:
  Cover metadata block       — system identity, life event, time horizon, audience
  Table of Contents          — Word TOC field (right-click → Update Field in Word)
  For each chapter (H1):
    Chapter narrative        — plain paragraph
    Quick reference rules    — H2 + bullet list
    For each worksheet (H2):
      Purpose / descriptor   — plain paragraph
      Estimated time         — italic metadata line
      Layout-specific body   — see layout dispatch below
      Decision gates         — H3 + pass/fail actions (all layouts)

Worksheet layout dispatch (DocxBuilder._add_worksheet):
  "form"       → _render_form()
                 Blue ALLCAPS section headers, field-type-aware blocks:
                   • select / radio  → radio-circle option list
                   • checkbox        → ☐ item
                   • text / default  → placeholder hint + grey fill-in underlines
  "table"      → _render_table()
                 Real Word table (Table Grid style): navy header row, alternating
                 row shading, one column per table_columns entry, row_count data rows.
  "checklist"  → _render_checklist()
                 Blue ☐ glyph + item text from checklist_items; formatted for print.
  "two-column" → _render_two_column()
                 3-column Word table: Field label / left_column_label / right_column_label,
                 navy header, alternating shading, one row per field across all sections.

Design choices:
  - Reads stage outputs from PipelineService.all_stage_outputs_as_dict() directly.
    No pipeline re-run or HTML render pass required.
  - chapter_expansion.chapters is preferred; falls back to worksheet_system.worksheets
    for pre-chapter-expansion projects.
  - Heading palette: H1 #1F569A (navy), H2 #2B6CB0 (blue), H3 #374151 (slate).
  - Writes to io.BytesIO — no filesystem writes.
"""
from __future__ import annotations

import io
from datetime import date
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.logging import get_logger

logger = get_logger(__name__)

_BLANK_LINE_COUNT = 2


def _add_toc_field(doc: Document) -> None:
    """Insert a Word TOC field at the current cursor position."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)

    update_instruction = doc.add_paragraph()
    update_instruction.add_run(
        '(Right-click this area in Word and select "Update Field" to generate the Table of Contents.)'
    ).italic = True
    update_instruction.runs[0].font.size = Pt(9)
    update_instruction.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def _add_page_break(doc: Document) -> None:
    """Add a page break paragraph."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _safe_text(value: Any, fallback: str = "") -> str:
    """Return a trimmed string from any value, or fallback if empty."""
    if not value:
        return fallback
    s = str(value).strip()
    return s if s else fallback


class DocxBuilder:
    """
    Accepts the dict of {stage_name: stage_output} returned by
    PipelineService.all_stage_outputs_as_dict() and produces a docx.Document
    serialised to bytes.
    """

    def build(self, project_id: int, all_outputs: dict[str, Any]) -> bytes:
        """
        Build the Word document and return it as raw bytes.

        Args:
            project_id:  Project ID (used for document_id formatting).
            all_outputs: Dict of stage_name → output dict for completed stages.

        Returns:
            Raw .docx bytes suitable for an HTTP attachment response.
        """
        doc = Document()
        self._set_default_styles(doc)

        arch = all_outputs.get("system_architecture", {})
        chapter_exp = all_outputs.get("chapter_expansion", {})
        ws_system = all_outputs.get("worksheet_system", {})

        system_name = _safe_text(arch.get("system_name"), "Operational Control System")
        life_event = _safe_text(arch.get("life_event"), "")
        system_objective = _safe_text(arch.get("system_objective"), "")
        time_horizon = _safe_text(arch.get("time_horizon"), "")
        audience = _safe_text(arch.get("audience"), "")
        document_id = f"LSB-{project_id:05d}"
        generated_date = date.today().strftime("%B %d, %Y")

        chapters: list[dict] = chapter_exp.get("chapters", [])
        legacy_worksheets: list[dict] = ws_system.get("worksheets", [])

        # ── Cover block ─────────────────────────────────────────────────────────
        doc.add_heading(system_name, level=0)

        cover_para = doc.add_paragraph()
        if life_event:
            cover_para.add_run(f"Life Event: {life_event}\n")
        if system_objective:
            cover_para.add_run(f"{system_objective}\n")
        if time_horizon:
            cover_para.add_run(f"Time Horizon: {time_horizon}\n")
        if audience:
            cover_para.add_run(f"Prepared For: {audience}\n")
        cover_para.add_run(f"Generated: {generated_date}  ·  Document ID: {document_id}")
        for run in cover_para.runs:
            run.font.size = Pt(10)

        doc.add_paragraph()

        # ── TOC field ────────────────────────────────────────────────────────────
        doc.add_heading("Table of Contents", level=1)
        _add_toc_field(doc)
        _add_page_break(doc)

        # ── Chapters ─────────────────────────────────────────────────────────────
        if chapters:
            total_worksheets = sum(len(ch.get("worksheets", [])) for ch in chapters)
            logger.debug(
                "DocxBuilder | project=%d | chapters=%d | worksheets=%d",
                project_id, len(chapters), total_worksheets,
            )
            for ch in chapters:
                ch_num = ch.get("chapter_number", chapters.index(ch) + 1)
                ch_title = _safe_text(ch.get("chapter_title"), f"Chapter {ch_num}")
                narrative = _safe_text(ch.get("narrative"), "")
                worksheets = ch.get("worksheets", [])

                doc.add_heading(f"Chapter {ch_num}: {ch_title}", level=1)

                if narrative:
                    doc.add_paragraph(narrative)

                quick_rules: list = ch.get("quick_reference_rules", [])
                if quick_rules:
                    doc.add_heading("Quick Reference Rules", level=2)
                    for rule in quick_rules:
                        rule_text = (
                            rule if isinstance(rule, str)
                            else rule.get("rule", rule.get("text", str(rule)))
                        )
                        if rule_text:
                            p = doc.add_paragraph(style="List Bullet")
                            p.add_run(_safe_text(rule_text))

                for ws in worksheets:
                    self._add_worksheet(doc, ws)

                _add_page_break(doc)

        elif legacy_worksheets:
            logger.debug(
                "DocxBuilder | project=%d | legacy path | worksheets=%d",
                project_id, len(legacy_worksheets),
            )
            for i, ws in enumerate(legacy_worksheets):
                title = _safe_text(ws.get("title"), f"Worksheet {i + 1}")
                purpose = _safe_text(ws.get("purpose"), "")
                doc.add_heading(title, level=1)
                if purpose:
                    doc.add_paragraph(purpose)
                self._add_worksheet(doc, ws)
                _add_page_break(doc)

        else:
            doc.add_paragraph(
                "No chapter content has been generated yet. "
                "Complete the Chapter Expansion pipeline stage to populate this document."
            )

        # ── Serialize to bytes ────────────────────────────────────────────────
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        raw = buf.read()
        logger.info(
            "DocxBuilder | project=%d | size=%d bytes",
            project_id, len(raw),
        )
        return raw

    # ── Layout-specific worksheet renderers ───────────────────────────────────

    def _add_worksheet(self, doc: Document, ws: dict) -> None:
        """
        Render a worksheet using its declared layout type.

        Dispatch table:
          "form"       → _render_form()        section headers + labelled fill-in fields
          "table"      → _render_table()       Word grid with column headers + data rows
          "checklist"  → _render_checklist()   ☐ checkbox items + decision gates
          "two-column" → _render_two_column()  side-by-side Word table (current / target)
          (unknown)    → _render_form() fallback
        """
        title   = _safe_text(ws.get("title"), "Worksheet")
        purpose = _safe_text(ws.get("purpose"), "")
        layout  = ws.get("layout", "form")
        est_time = _safe_text(ws.get("estimated_completion_time"), "")

        doc.add_heading(title, level=2)

        if purpose:
            p = doc.add_paragraph(purpose)
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        if est_time:
            meta = doc.add_paragraph()
            r = meta.add_run(f"Estimated completion time: {est_time}")
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        if layout == "table":
            self._render_table(doc, ws)
        elif layout == "checklist":
            self._render_checklist(doc, ws)
        elif layout == "two-column":
            self._render_two_column(doc, ws)
        else:
            self._render_form(doc, ws)

        # Decision gates (shared across layouts)
        gates: list[dict] = ws.get("decision_gates", [])
        if gates:
            doc.add_heading("Decision Gates", level=3)
            for gate in gates:
                cond = _safe_text(gate.get("condition") or gate.get("gate_title"), "")
                pass_a = _safe_text(gate.get("pass_action"), "")
                fail_a = _safe_text(gate.get("fail_action"), "")
                if cond:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(f"Gate: {cond}").bold = True
                    if pass_a:
                        doc.add_paragraph(f"  ✓ Pass → {pass_a}")
                    if fail_a:
                        doc.add_paragraph(f"  ✗ Fail → {fail_a}")

    def _render_form(self, doc: Document, ws: dict) -> None:
        """
        Form layout — sections with bold section headers, each field as a
        labelled block with a fill-in underline beneath it.
        """
        sections: list[dict] = ws.get("sections", [])

        # Legacy fallback: worksheets without sections but with top-level fields
        if not sections:
            for field_item in ws.get("fields", []):
                label = _safe_text(
                    field_item.get("label") or field_item.get("name") or field_item.get("title"), ""
                )
                if label:
                    doc.add_heading(label, level=3)
                    _add_fill_line(doc)
            if not ws.get("fields"):
                doc.add_heading("Notes", level=3)
                for _ in range(3):
                    _add_fill_line(doc)
            return

        for section in sections:
            sec_title = _safe_text(section.get("section_title"), "")
            instructions = _safe_text(section.get("instructions"), "")

            if sec_title:
                sec_p = doc.add_paragraph()
                r = sec_p.add_run(sec_title.upper())
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x1F, 0x56, 0x9A)

            if instructions:
                ins_p = doc.add_paragraph(instructions)
                for run in ins_p.runs:
                    run.font.size = Pt(9)
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            fields: list[dict] = section.get("fields", [])
            if fields:
                for field_item in fields:
                    label = _safe_text(
                        field_item.get("label") or field_item.get("name") or field_item.get("title"), ""
                    )
                    ftype = _safe_text(field_item.get("type"), "text")
                    placeholder = _safe_text(field_item.get("placeholder"), "")
                    if not label:
                        continue

                    doc.add_heading(label, level=3)

                    if ftype in ("select", "radio") and field_item.get("options"):
                        # Multiple choice — render each option as a selectable item
                        for opt in field_item["options"]:
                            opt_text = _safe_text(opt, "")
                            if opt_text:
                                p = doc.add_paragraph()
                                p.add_run("○  ").font.size = Pt(11)
                                p.add_run(opt_text).font.size = Pt(10)
                    elif ftype == "checkbox":
                        p = doc.add_paragraph()
                        p.add_run("☐  ").font.size = Pt(11)
                        p.add_run(placeholder or label).font.size = Pt(10)
                    else:
                        if placeholder:
                            ph_p = doc.add_paragraph(f"({placeholder})")
                            for run in ph_p.runs:
                                run.font.size = Pt(9)
                                run.italic = True
                                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                        _add_fill_line(doc)
                        _add_fill_line(doc)
            else:
                doc.add_heading("Notes", level=3)
                for _ in range(3):
                    _add_fill_line(doc)

    def _render_table(self, doc: Document, ws: dict) -> None:
        """
        Table layout — a Word table with bold column headers and blank data rows.
        """
        columns: list[str] = ws.get("table_columns", [])
        row_count: int = min(ws.get("table_row_count", 12), 20)

        if not columns:
            doc.add_heading("Notes", level=3)
            _add_fill_line(doc)
            return

        col_count = len(columns)
        table = doc.add_table(rows=row_count + 1, cols=col_count)
        table.style = "Table Grid"

        # Header row
        hdr_row = table.rows[0]
        for i, col_name in enumerate(columns):
            cell = hdr_row.cells[i]
            cell.text = col_name
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # Navy header shading
            _shade_cell(cell, "1F569A")
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Data rows — blank, with alternating light shading
        for row_idx in range(1, row_count + 1):
            for cell in table.rows[row_idx].cells:
                cell.text = ""
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
                if row_idx % 2 == 0:
                    _shade_cell(cell, "EEF3FA")

        doc.add_paragraph()  # spacing after table

    def _render_checklist(self, doc: Document, ws: dict) -> None:
        """
        Checklist layout — ☐ checkbox symbol + item text, styled for print use.
        Decision gates follow (handled by caller).
        """
        items: list = ws.get("checklist_items", [])

        if not items:
            doc.add_heading("Checklist Items", level=3)
            for _ in range(5):
                p = doc.add_paragraph()
                p.add_run("☐  _______________________________________________")
                p.runs[0].font.size = Pt(11)
            return

        for item in items:
            item_text = _safe_text(
                item if isinstance(item, str) else (item.get("label") or item.get("text") or str(item)), ""
            )
            if not item_text:
                continue
            p = doc.add_paragraph()
            chk = p.add_run("☐  ")
            chk.font.size = Pt(12)
            chk.font.color.rgb = RGBColor(0x1F, 0x56, 0x9A)
            body = p.add_run(item_text)
            body.font.size = Pt(10)

    def _render_two_column(self, doc: Document, ws: dict) -> None:
        """
        Two-column layout — a 2-column Word table.
        Left header = left_column_label, right header = right_column_label.
        Each field generates one row: label in first column, fill space in second.
        """
        left_label  = _safe_text(ws.get("left_column_label"),  "Current State")
        right_label = _safe_text(ws.get("right_column_label"), "Target State")
        sections: list[dict] = ws.get("sections", [])

        # Collect all fields across sections preserving order
        rows_data: list[tuple[str, str]] = []  # (section_header | "", field_label)
        for section in sections:
            sec_title = _safe_text(section.get("section_title"), "")
            for i, field_item in enumerate(section.get("fields", [])):
                label = _safe_text(
                    field_item.get("label") or field_item.get("name") or field_item.get("title"), ""
                )
                if not label:
                    continue
                rows_data.append((sec_title if i == 0 else "", label))

        if not rows_data:
            rows_data = [("", "Notes")]

        # Build table: 1 header row + 1 row per field
        table = doc.add_table(rows=len(rows_data) + 1, cols=3)
        table.style = "Table Grid"

        # Column widths via XML (approx): label ~30%, left ~35%, right ~35%
        col_widths = ["2200", "2500", "2500"]  # twips
        for i, cell in enumerate(table.rows[0].cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), col_widths[i])
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)

        # Header row
        headers = ["Field", left_label, right_label]
        for i, hdr_text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = hdr_text
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade_cell(cell, "1F569A")
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Data rows
        for row_idx, (sec_title, field_label) in enumerate(rows_data):
            row = table.rows[row_idx + 1]
            label_text = f"{sec_title} — {field_label}" if sec_title else field_label
            row.cells[0].text = label_text
            for run in row.cells[0].paragraphs[0].runs:
                run.font.size = Pt(9)
                run.bold = bool(sec_title)
            row.cells[1].text = ""
            row.cells[2].text = ""
            if row_idx % 2 == 0:
                _shade_cell(row.cells[1], "EEF3FA")
                _shade_cell(row.cells[2], "EEF3FA")

        doc.add_paragraph()  # spacing after table

    def _set_default_styles(self, doc: Document) -> None:
        """Adjust base font and heading colours for a clean, professional look."""
        style = doc.styles["Normal"]
        style.font.name  = "Calibri"
        style.font.size  = Pt(11)

        # Heading colour palette: navy blue for H1/H2, slate for H3
        _heading_colors = {
            "Heading 1": ("1F569A", 16),
            "Heading 2": ("2B6CB0", 13),
            "Heading 3": ("374151", 11),
        }
        for style_name, (hex_color, pt_size) in _heading_colors.items():
            try:
                h = doc.styles[style_name]
                h.font.name  = "Calibri"
                h.font.size  = Pt(pt_size)
                r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
                h.font.color.rgb = RGBColor(r, g, b)
            except KeyError:
                pass


# ── Module-level helpers ──────────────────────────────────────────────────────

def _add_fill_line(doc: Document) -> None:
    """One printable underline fill-in line."""
    p = doc.add_paragraph()
    r = p.add_run("_" * 72)
    r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    r.font.size = Pt(10)


def _shade_cell(cell: Any, hex_color: str) -> None:
    """Apply a solid background colour to a table cell via OOXML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)
